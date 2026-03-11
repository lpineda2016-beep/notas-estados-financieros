import pandas as pd
from docx import Document


def generar_notas(excel_file):

    # Leer Excel
    balance = pd.read_excel(excel_file)

    # Normalizar nombres de columnas
    balance.columns = balance.columns.str.strip().str.lower()

    # Validar columnas requeridas
    columnas = ["cuenta", "nota", "valor_2025", "valor_2024"]

    for col in columnas:
        if col not in balance.columns:
            raise ValueError(f"Falta la columna: {col}")

    # Agrupar por nota
    notas = balance.groupby("nota")[["valor_2025", "valor_2024"]].sum()

    # Crear documento Word
    doc = Document()

    doc.add_heading("NOTAS A LOS ESTADOS FINANCIEROS", 0)
    doc.add_paragraph("Al 31 de diciembre de 2025")
    doc.add_paragraph("Valores expresados en USD")

    titulos_notas = {
        4: "Efectivo y equivalentes de efectivo",
        5: "Cuentas por cobrar",
        6: "Inventarios",
        7: "Propiedad planta y equipo",
        8: "Cuentas por pagar",
        10: "Patrimonio",
        11: "Ingresos",
        12: "Gastos operacionales"
    }

    for nota, valores in notas.iterrows():

        titulo = titulos_notas.get(nota, f"Nota {nota}")

        doc.add_heading(f"Nota {nota} - {titulo}", level=1)

        table = doc.add_table(rows=1, cols=3)

        headers = table.rows[0].cells
        headers[0].text = "Concepto"
        headers[1].text = "2025"
        headers[2].text = "2024"

        cuentas = balance[balance["nota"] == nota]

        for _, row in cuentas.iterrows():

            cells = table.add_row().cells
            cells[0].text = str(row["cuenta"])
            cells[1].text = str(row["valor_2025"])
            cells[2].text = str(row["valor_2024"])

        total = table.add_row().cells
        total[0].text = "TOTAL"
        total[1].text = str(valores["valor_2025"])
        total[2].text = str(valores["valor_2024"])

    output = "notas_estados_financieros.docx"
    doc.save(output)

    return output
